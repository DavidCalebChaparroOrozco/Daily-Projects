# Importing necessary libraries
from docx import Document
from docx.shared import Pt
import matplotlib.pyplot as plt
import io
import pandas as pd

class DocumentView:
    def __init__(self, filename="report.docx"):
        self.filename = filename
        self.document = Document()

    # Add a paragraph to the document
    def add_paragraph(self, text, style="Arial", size=12, bold=False):
        paragraph = self.document.add_paragraph(text)
        run = paragraph.runs[0]
        run.font.name = style
        run.font.size = Pt(size)
        run.bold = bold

    # Add a table to the document
    def add_table(self, data):
        table = self.document.add_table(rows=1, cols=len(data.columns))
        hdr_cells = table.rows[0].cells
        for i, column in enumerate(data.columns):
            hdr_cells[i].text = column
        for index, row in data.iterrows():
            row_cells = table.add_row().cells
            for i, column in enumerate(data.columns):
                row_cells[i].text = str(row[column])

    # Add an image to the document
    def add_image(self, image_stream):
        self.document.add_picture(image_stream, width=Pt(400))

    # Save the document
    def save(self):
        self.document.save(self.filename)
        print(f"Document {self.filename} saved successfully")

class GraphView:
    @staticmethod
    # Plot hours worked over time (Line Chart).
    def plot_hours_worked_over_time(data):
        plt.figure(figsize=(10, 5))
        data['Date'] = pd.to_datetime(data['Date'])
        data.plot(x='Date', y='Hours Worked', kind='line', title='Hours Worked Over Time')
        buf = io.BytesIO()
        plt.savefig(buf, format='png')
        buf.seek(0)
        plt.close()
        return buf

    @staticmethod
    # Plot total hours worked by each employee (Bar Chart).
    def plot_total_hours_by_employee(data):
        plt.figure(figsize=(10, 5))
        total_hours = data.groupby('Employee')['Hours Worked'].sum()
        total_hours.plot(kind='bar', title='Total Hours Worked by Employee')
        buf = io.BytesIO()
        plt.savefig(buf, format='png')
        buf.seek(0)
        plt.close()
        return buf

    @staticmethod
    # Plot distribution of hours worked by client (Pie Chart).
    def plot_hours_distribution_by_client(data):
        plt.figure(figsize=(8, 8))
        client_hours = data.groupby('Client')['Hours Worked'].sum()
        client_hours.plot(kind='pie', autopct='%1.1f%%', title='Hours Worked by Client')
        buf = io.BytesIO()
        plt.savefig(buf, format='png')
        buf.seek(0)
        plt.close()
        return buf

    @staticmethod
    # Plot histogram of hours worked (Histogram).
    def plot_hours_histogram(data):
        plt.figure(figsize=(10, 5))
        data['Hours Worked'].plot(kind='hist', bins=10, title='Distribution of Hours Worked')
        buf = io.BytesIO()
        plt.savefig(buf, format='png')
        buf.seek(0)
        plt.close()
        return buf

    @staticmethod
    # Plot scatter plot of hours worked vs date (Scatter Plot).
    def plot_scatter_hours_vs_date(data):
        plt.figure(figsize=(10, 5))
        data['Date'] = pd.to_datetime(data['Date'])
        plt.scatter(data['Date'], data['Hours Worked'])
        plt.title('Hours Worked vs Date')
        plt.xlabel('Date')
        plt.ylabel('Hours Worked')
        buf = io.BytesIO()
        plt.savefig(buf, format='png')
        buf.seek(0)
        plt.close()
        return buf